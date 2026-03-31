from docx import Document
from docx.shared import Inches
import matplotlib.pyplot as plt
import numpy as np
from io import BytesIO

# Generate the plot
t = np.linspace(0.0, 2.0, 100)
p = np.sin(2 * np.pi * t)
plt.plot(t, p)
plt.xlabel("Time (t)")
plt.ylabel("Pressure (p)")
plt.title("Pressure vs. Time")
plt.grid(True)

# Save the plot to a memory file
memfile = BytesIO()
plt.savefig(memfile)
plt.show()
plt.close()  # Close the plot to free memory

# Create a Word document
document = Document()
document.add_heading("Pressure vs. Time Report", level=0)
document.add_picture(memfile, width=Inches(4))  # Adjust width as needed
document.save("pressure_report.docx")
memfile.close()  # Close the memory file

print("Graph saved in pressure_report.docx")
